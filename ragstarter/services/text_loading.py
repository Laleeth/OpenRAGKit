from __future__ import annotations

import csv
import hashlib
import mimetypes
from dataclasses import dataclass
from io import StringIO
from pathlib import Path
from urllib.parse import urlparse

import httpx
import fitz  # PyMuPDF
from bs4 import BeautifulSoup
from docx import Document as DocxDocument


@dataclass(slots=True)
class LoadedDocument:
    text: str
    source_name: str
    source_type: str
    source_uri: str | None
    title: str | None
    content_hash: str
    text_length: int


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()


def guess_source_type(path: Path | str) -> str:
    if isinstance(path, Path):
        ext = path.suffix.lower()
        if ext == ".pdf":
            return "pdf"
        if ext in {".md", ".markdown"}:
            return "markdown"
        if ext == ".txt":
            return "text"
        if ext == ".docx":
            return "docx"
        if ext in {".html", ".htm"}:
            return "html"
        if ext == ".csv":
            return "csv"
        return ext.removeprefix(".") or "file"
    parsed = urlparse(str(path))
    if parsed.scheme in {"http", "https"}:
        return "url"
    return "text"


def load_from_url(url: str, timeout_s: float = 30.0) -> LoadedDocument:
    response = httpx.get(url, timeout=timeout_s, follow_redirects=True)
    response.raise_for_status()
    content_type = response.headers.get("content-type", "").lower()
    if "pdf" in content_type or url.lower().endswith(".pdf"):
        return load_pdf_bytes(response.content, source_name=url, source_uri=url)
    if "html" in content_type or "<html" in response.text.lower():
        return load_html(response.text, source_name=url, source_uri=url)
    return load_text(
        response.text,
        source_name=url,
        source_type="url",
        source_uri=url,
        title=url,
    )


def load_file(path: Path) -> LoadedDocument:
    source_type = guess_source_type(path)
    if source_type == "pdf":
        return load_pdf(path)
    if source_type == "docx":
        return load_docx(path)
    if source_type in {"html", "htm"}:
        return load_html(path.read_text(encoding="utf-8", errors="ignore"), source_name=path.name, source_uri=str(path))
    if source_type == "csv":
        return load_csv(path)
    return load_text(path.read_text(encoding="utf-8", errors="ignore"), source_name=path.name, source_type=source_type, source_uri=str(path), title=path.stem)


def load_pdf(path: Path) -> LoadedDocument:
    return load_pdf_bytes(path.read_bytes(), source_name=path.name, source_uri=str(path))


def load_pdf_bytes(data: bytes, source_name: str, source_uri: str | None = None) -> LoadedDocument:
    doc = fitz.open(stream=data, filetype="pdf")
    parts: list[str] = []
    for page_number, page in enumerate(doc, start=1):
        text = page.get_text("text").strip()
        if text:
            parts.append(f"[Page {page_number}]\n{text}")
    text = "\n\n".join(parts).strip()
    title = doc.metadata.get("title") if doc.metadata else None
    content_hash = sha256_text(text or source_name)
    return LoadedDocument(
        text=text,
        source_name=source_name,
        source_type="pdf",
        source_uri=source_uri,
        title=title or Path(source_name).stem,
        content_hash=content_hash,
        text_length=len(text),
    )


def load_docx(path: Path) -> LoadedDocument:
    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n\n".join(paragraphs)
    title = path.stem
    try:
        if doc.core_properties.title:
            title = doc.core_properties.title
    except Exception:
        pass
    return LoadedDocument(
        text=text,
        source_name=path.name,
        source_type="docx",
        source_uri=str(path),
        title=title,
        content_hash=sha256_text(text or path.name),
        text_length=len(text),
    )


def load_html(html: str, source_name: str, source_uri: str | None = None) -> LoadedDocument:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    title = None
    if soup.title and soup.title.string:
        title = soup.title.string.strip()
    text = soup.get_text("\n")
    lines = [line.strip() for line in text.splitlines()]
    cleaned = "\n".join(line for line in lines if line)
    return LoadedDocument(
        text=cleaned,
        source_name=source_name,
        source_type="html",
        source_uri=source_uri,
        title=title or Path(source_name).stem,
        content_hash=sha256_text(cleaned or source_name),
        text_length=len(cleaned),
    )


def load_csv(path: Path) -> LoadedDocument:
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        rows = list(reader)
    if not rows:
        text = ""
    else:
        headers = rows[0]
        body = rows[1:]
        lines = [" | ".join(headers), " | ".join(["---"] * len(headers))]
        for row in body[:5000]:
            row = list(row)[: len(headers)]
            row += [""] * (len(headers) - len(row))
            lines.append(" | ".join(row))
        text = "\n".join(lines)
    return LoadedDocument(
        text=text,
        source_name=path.name,
        source_type="csv",
        source_uri=str(path),
        title=path.stem,
        content_hash=sha256_text(text or path.name),
        text_length=len(text),
    )


def load_text(text: str, source_name: str, source_type: str, source_uri: str | None = None, title: str | None = None) -> LoadedDocument:
    cleaned = "\n".join(line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"))
    cleaned = cleaned.strip()
    return LoadedDocument(
        text=cleaned,
        source_name=source_name,
        source_type=source_type,
        source_uri=source_uri,
        title=title or Path(source_name).stem,
        content_hash=sha256_text(cleaned or source_name),
        text_length=len(cleaned),
    )
